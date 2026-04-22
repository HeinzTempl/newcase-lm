"""
Kanzlei Pipeline v2 - DOCX Export
==================================
Konvertiert die Pipeline-Outputs (Markdown) in formatierte Word-Dokumente.
"""

import logging
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


def _setup_styles(doc: Document):
    """Konfiguriert die Basis-Styles des Dokuments."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def _add_header(doc: Document, title: str, subtitle: str, confidential: bool = False):
    """Fügt Titel, Datum und optionalen Vertraulichkeitsvermerk hinzu."""
    # Titel
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Datum
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(subtitle)
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    # Vertraulichkeitsvermerk
    if confidential:
        conf_para = doc.add_paragraph()
        conf_run = conf_para.add_run("VERTRAULICH – NUR FÜR KANZLEIINTERNEN GEBRAUCH")
        conf_run.bold = True
        conf_run.font.size = Pt(10)
        conf_run.font.color.rgb = RGBColor(180, 0, 0)

    # Trennlinie
    doc.add_paragraph("─" * 60)


def _add_doc_overview_table(doc: Document, overview_text: str):
    """Parst die Markdown-Tabelle und erstellt eine DOCX-Tabelle."""
    lines = overview_text.strip().split("\n")

    # Überschrift der Übersicht
    for line in lines:
        if line.startswith("##"):
            doc.add_heading(line.lstrip("#").strip(), level=2)
            break

    # Tabelle parsen
    table_lines = [l for l in lines if l.startswith("|")]
    if len(table_lines) < 3:
        # Keine gültige Tabelle gefunden, als Text einfügen
        doc.add_paragraph(overview_text)
        return

    # Header-Zeile parsen
    headers = [cell.strip() for cell in table_lines[0].split("|")[1:-1]]
    # Separator überspringen (table_lines[1])
    data_rows = table_lines[2:]

    # Tabelle erstellen
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Daten
    for row_text in data_rows:
        cells = [cell.strip() for cell in row_text.split("|")[1:-1]]
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            if i < len(row.cells):
                row.cells[i].text = cell_text
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    doc.add_paragraph()  # Abstand nach Tabelle


def _parse_markdown_to_docx(doc: Document, md_text: str):
    """
    Parst Markdown-Text und schreibt ihn formatiert ins DOCX.
    Unterstützt: Überschriften, Fettdruck, Aufzählungen, Trennlinien.
    """
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Leerzeile
        if not stripped:
            i += 1
            continue

        # Trennlinie
        if stripped.startswith("---") or stripped.startswith("───"):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # Überschriften
        if stripped.startswith("# "):
            doc.add_heading(stripped.lstrip("#").strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped.lstrip("#").strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped.lstrip("#").strip(), level=3)
            i += 1
            continue

        # Aufzählungspunkte
        if stripped.startswith("* ") or stripped.startswith("- "):
            bullet_text = stripped[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, bullet_text)
            i += 1
            continue

        # Tabelle überspringen (wird separat behandelt)
        if stripped.startswith("|"):
            i += 1
            continue

        # Kursiv-nur Zeile (z.B. *Erstellt: ...*)
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            para = doc.add_paragraph()
            run = para.add_run(stripped.strip("*"))
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Normaler Absatz - sammle zusammenhängende Zeilen
        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if (not next_line or next_line.startswith("#") or
                next_line.startswith("---") or next_line.startswith("|") or
                next_line.startswith("* ") or next_line.startswith("- ")):
                break
            paragraph_lines.append(next_line)
            i += 1

        full_text = " ".join(paragraph_lines)
        para = doc.add_paragraph()
        _add_formatted_text(para, full_text)


def _add_formatted_text(paragraph, text: str):
    """Fügt Text mit Markdown-Fettdruck (**text**) formatiert ein."""
    # Splitte an **...**
    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def export_klartext_docx(
    act_summary: str,
    doc_overview: str,
    output_path: Path,
    timestamp: str = None,
    mapping_table: str = None,
):
    """
    Exportiert die Klartext-Gesamtübersicht als DOCX.

    Args:
        act_summary: Der LLM-generierte Sachverhalt (Markdown)
        doc_overview: Die Dokumentenübersicht-Tabelle (Markdown)
        output_path: Zielpfad für die DOCX-Datei
        mapping_table: Optionale Zuordnungstabelle (Klartext → Anonymisiert)
    """
    if not timestamp:
        timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")

    doc = Document()
    _setup_styles(doc)

    _add_header(
        doc,
        title="Gesamtübersicht Akt",
        subtitle=f"Erstellt: {timestamp}",
        confidential=True,
    )

    # Dokumentenübersicht als Tabelle
    _add_doc_overview_table(doc, doc_overview)

    # Trennlinie
    doc.add_paragraph("─" * 60)

    # Sachverhalt
    _parse_markdown_to_docx(doc, act_summary)

    # Zuordnungstabelle (Klartext → Anonymisiert)
    if mapping_table:
        doc.add_paragraph("─" * 60)
        doc.add_heading("Zuordnung Klartext → Anonymisiert", level=2)
        _add_doc_overview_table(doc, mapping_table)

    # Speichern
    doc.save(str(output_path))
    logger.info(f"  → DOCX: {output_path.name}")


def export_anon_docx(
    anon_summary: str,
    output_path: Path,
    timestamp: str = None,
    doc_overview: str = None,
):
    """
    Exportiert die anonymisierte Gesamtübersicht als DOCX.

    Args:
        anon_summary: Der anonymisierte Sachverhalt (Markdown)
        output_path: Zielpfad für die DOCX-Datei
        doc_overview: Optionale anonymisierte Dokumentenübersicht (Markdown)
    """
    if not timestamp:
        timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")

    doc = Document()
    _setup_styles(doc)

    _add_header(
        doc,
        title="Gesamtübersicht Akt (Anonymisiert)",
        subtitle=f"Erstellt: {timestamp} · Anonymisiert für Cloud-LLM-Nutzung",
        confidential=False,
    )

    # Anonymisierte Dokumentenübersicht
    if doc_overview:
        _add_doc_overview_table(doc, doc_overview)
        doc.add_paragraph("─" * 60)

    # Sachverhalt
    _parse_markdown_to_docx(doc, anon_summary)

    # Speichern
    doc.save(str(output_path))
    logger.info(f"  → DOCX: {output_path.name}")
